from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT_PATH = r"C:\Users\SBS\Desktop\agent02\docx_creator-workspace\iteration-1\eval-2\without_skill\outputs\resume.docx"

doc = Document()

# ── 페이지 여백 설정 ──────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── 헬퍼 함수 ────────────────────────────────────────────
def set_font(run, name="맑은 고딕", size=11, bold=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    # 한글 폰트 지정
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:eastAsia"), name)
    rPr.insert(0, rFonts)

def add_horizontal_rule(doc):
    """가로 구분선 추가"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "2E4057")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def add_section_title(doc, title):
    """섹션 제목 스타일"""
    add_horizontal_rule(doc)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(title)
    set_font(run, size=13, bold=True, color=(46, 64, 87))
    return p

def shade_cell(cell, fill_hex):
    """테이블 셀 배경색"""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  fill_hex)
    tcPr.append(shd)

def set_cell_border(table):
    """테이블 전체 테두리 설정"""
    for row in table.rows:
        for cell in row.cells:
            tc   = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBdr = OxmlElement("w:tcBdr")
            for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
                el = OxmlElement(f"w:{side}")
                el.set(qn("w:val"),   "single")
                el.set(qn("w:sz"),    "4")
                el.set(qn("w:space"), "0")
                el.set(qn("w:color"), "CCCCCC")
                tcBdr.append(el)
            tcPr.append(tcBdr)

# ════════════════════════════════════════════════════════
# 1. 이름 / 직함
# ════════════════════════════════════════════════════════
p_name = doc.add_paragraph()
p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_name.paragraph_format.space_before = Pt(0)
p_name.paragraph_format.space_after  = Pt(4)
run_name = p_name.add_run("홍  길  동")
set_font(run_name, size=26, bold=True, color=(46, 64, 87))

p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_title.paragraph_format.space_after = Pt(2)
run_title = p_title.add_run("소프트웨어 엔지니어")
set_font(run_title, size=12, color=(100, 110, 120))

# ── 구분선 ──
add_horizontal_rule(doc)

# ════════════════════════════════════════════════════════
# 2. 연락처 정보
# ════════════════════════════════════════════════════════
contact_items = [
    ("전화", "010-1234-5678"),
    ("이메일", "hong.gildong@email.com"),
    ("주소", "서울특별시 강남구 테헤란로 123"),
    ("GitHub", "github.com/hong-gildong"),
    ("LinkedIn", "linkedin.com/in/hong-gildong"),
]

p_contact = doc.add_paragraph()
p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_contact.paragraph_format.space_before = Pt(4)
p_contact.paragraph_format.space_after  = Pt(2)

contact_line = "  |  ".join(f"{k}: {v}" for k, v in contact_items[:3])
run_c = p_contact.add_run(contact_line)
set_font(run_c, size=9, color=(80, 90, 100))

p_contact2 = doc.add_paragraph()
p_contact2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_contact2.paragraph_format.space_after = Pt(6)
contact_line2 = "  |  ".join(f"{k}: {v}" for k, v in contact_items[3:])
run_c2 = p_contact2.add_run(contact_line2)
set_font(run_c2, size=9, color=(80, 90, 100))

# ════════════════════════════════════════════════════════
# 3. 자기소개
# ════════════════════════════════════════════════════════
add_section_title(doc, "자기소개")

intro_text = (
    "5년 이상의 백엔드 개발 경험을 보유한 소프트웨어 엔지니어입니다. "
    "Python 및 Java 기반 웹 서비스 설계·개발과 클라우드 인프라 운영에 강점이 있으며, "
    "대규모 트래픽 환경에서의 성능 최적화 및 마이크로서비스 아키텍처 전환 프로젝트를 성공적으로 이끈 경험이 있습니다. "
    "새로운 기술에 대한 적극적인 학습 태도와 팀 내 원활한 협업 능력을 바탕으로, "
    "비즈니스 목표와 기술적 품질을 함께 달성하는 개발자를 지향합니다."
)
p_intro = doc.add_paragraph()
p_intro.paragraph_format.space_after = Pt(4)
run_intro = p_intro.add_run(intro_text)
set_font(run_intro, size=10, color=(50, 50, 50))
p_intro.paragraph_format.first_line_indent = Pt(0)

# ════════════════════════════════════════════════════════
# 4. 경력사항 표
# ════════════════════════════════════════════════════════
add_section_title(doc, "경력사항")

career_data = [
    ("회사명",           "근무 기간",                  "담당 업무"),  # 헤더
    ("(주)테크스타트",   "2022.03 – 현재",
     "• Python/FastAPI 기반 REST API 설계 및 개발\n"
     "• AWS ECS·RDS 기반 서비스 인프라 구축 및 운영\n"
     "• CI/CD 파이프라인 구성(GitHub Actions, ArgoCD)"),
    ("네오소프트",       "2020.01 – 2022.02",
     "• Spring Boot 기반 전자상거래 플랫폼 백엔드 개발\n"
     "• 결제 모듈 연동(PG사 API) 및 보안 취약점 개선\n"
     "• 월간 활성 사용자 50만 규모 서비스 성능 튜닝"),
    ("디지털웍스",       "2018.07 – 2019.12",
     "• Java/JSP 기반 사내 ERP 시스템 유지보수\n"
     "• MySQL 쿼리 최적화 및 배치 처리 스크립트 작성\n"
     "• 신규 입사자 온보딩 문서 및 API 명세 작성"),
]

col_widths = [Cm(4.0), Cm(3.5), Cm(10.5)]
tbl = doc.add_table(rows=len(career_data), cols=3)
tbl.style = "Table Grid"

for r_idx, row_data in enumerate(career_data):
    row = tbl.rows[r_idx]
    for c_idx, cell_text in enumerate(row_data):
        cell = row.cells[c_idx]
        cell.width = col_widths[c_idx]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(4)
        p.paragraph_format.left_indent  = Pt(4)

        is_header = r_idx == 0
        run = p.add_run(cell_text)
        set_font(run, size=9 if not is_header else 10,
                 bold=is_header,
                 color=(255, 255, 255) if is_header else (40, 40, 40))

        if is_header:
            shade_cell(cell, "2E4057")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif r_idx % 2 == 0:
            shade_cell(cell, "F2F5F8")

set_cell_border(tbl)

# ════════════════════════════════════════════════════════
# 5. 학력
# ════════════════════════════════════════════════════════
add_section_title(doc, "학력")

edu_data = [
    ("한국대학교",   "컴퓨터공학과 학사",  "2014.03 – 2018.02",  "졸업"),
    ("서울고등학교", "이과",               "2011.03 – 2014.02",  "졸업"),
]

edu_tbl = doc.add_table(rows=len(edu_data) + 1, cols=4)
edu_tbl.style = "Table Grid"

headers = ["학교명", "전공 / 계열", "재학 기간", "비고"]
edu_col_widths = [Cm(4.5), Cm(4.5), Cm(4.0), Cm(2.5)]

header_row = edu_tbl.rows[0]
for c_idx, h in enumerate(headers):
    cell = header_row.cells[c_idx]
    cell.width = edu_col_widths[c_idx]
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(h)
    set_font(run, size=10, bold=True, color=(255, 255, 255))
    shade_cell(cell, "2E4057")

for r_idx, row_data in enumerate(edu_data):
    row = edu_tbl.rows[r_idx + 1]
    for c_idx, cell_text in enumerate(row_data):
        cell = row.cells[c_idx]
        cell.width = edu_col_widths[c_idx]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        run = p.add_run(cell_text)
        fill = "F2F5F8" if r_idx % 2 == 0 else "FFFFFF"
        set_font(run, size=9, color=(40, 40, 40))
        shade_cell(cell, fill)

set_cell_border(edu_tbl)

# ════════════════════════════════════════════════════════
# 6. 보유 기술
# ════════════════════════════════════════════════════════
add_section_title(doc, "보유 기술")

skill_categories = [
    ("프로그래밍 언어",
     ["Python", "Java", "JavaScript (ES6+)", "SQL", "Bash/Shell"]),
    ("프레임워크 / 라이브러리",
     ["FastAPI", "Spring Boot", "Django", "React", "Pandas / NumPy"]),
    ("데이터베이스",
     ["PostgreSQL", "MySQL", "Redis", "MongoDB", "Elasticsearch"]),
    ("클라우드 / 인프라",
     ["AWS (EC2, ECS, RDS, S3, Lambda)", "Docker", "Kubernetes", "Terraform", "GitHub Actions"]),
    ("도구 / 협업",
     ["Git / GitHub", "Jira", "Confluence", "Postman", "VS Code / IntelliJ"]),
    ("자격증",
     ["AWS Certified Solutions Architect – Associate (2023)",
      "정보처리기사 (2018)",
      "SQLD (2019)"]),
]

for category, skills in skill_categories:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(1)

    run_cat = p.add_run(f"{category}:  ")
    set_font(run_cat, size=10, bold=True, color=(46, 64, 87))

    run_skills = p.add_run("  ·  ".join(skills))
    set_font(run_skills, size=9, color=(50, 50, 50))

# ── 하단 여백 ──────────────────────────────────────────
doc.add_paragraph()

# ── 저장 ──────────────────────────────────────────────
doc.save(OUTPUT_PATH)
print(f"resume.docx 저장 완료: {OUTPUT_PATH}")
