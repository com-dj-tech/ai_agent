from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_PATH = "C:/Users/SBS/Desktop/agent02/docx_creator-workspace/iteration-1/eval-2/with_skill/outputs/resume.docx"


def set_korean_font(run, font_name="맑은 고딕", size=None, bold=False, color=None):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    if size:
        run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color


def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def add_section_heading(doc, title):
    """섹션 구분선 + 제목 추가"""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(14)
    para.paragraph_format.space_after = Pt(4)
    run = para.add_run(title)
    set_korean_font(run, size=13, bold=True, color=RGBColor(0x1F, 0x49, 0x7D))
    # 하단 테두리(구분선) 적용
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1F497D")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return para


def main():
    doc = Document()

    # ── 페이지 설정 (A4, 여백) ──────────────────────────────────────────────
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.8)

    # ── 이름 (대제목) ──────────────────────────────────────────────────────
    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_para.paragraph_format.space_after = Pt(2)
    name_run = name_para.add_run("홍  길  동")
    set_korean_font(name_run, size=28, bold=True, color=RGBColor(0x1F, 0x49, 0x7D))

    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_para.paragraph_format.space_after = Pt(10)
    sub_run = subtitle_para.add_run("소프트웨어 엔지니어")
    set_korean_font(sub_run, size=13, color=RGBColor(0x59, 0x59, 0x59))

    # ── 연락처 정보 ────────────────────────────────────────────────────────
    contact_para = doc.add_paragraph()
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_para.paragraph_format.space_after = Pt(14)
    contact_items = [
        ("전화", "010-1234-5678"),
        ("이메일", "hong.gildong@email.com"),
        ("GitHub", "github.com/honggildong"),
        ("주소", "서울특별시 강남구"),
    ]
    contact_text = "   |   ".join(f"{k}: {v}" for k, v in contact_items)
    c_run = contact_para.add_run(contact_text)
    set_korean_font(c_run, size=10, color=RGBColor(0x44, 0x44, 0x44))

    # ── 자기소개 ───────────────────────────────────────────────────────────
    add_section_heading(doc, "자기소개")
    intro_para = doc.add_paragraph()
    intro_para.paragraph_format.space_after = Pt(4)
    intro_para.paragraph_format.line_spacing = Pt(20)
    intro_text = (
        "5년 이상의 백엔드 개발 경험을 보유한 소프트웨어 엔지니어로, "
        "Python 및 Java 기반의 서버 개발과 클라우드 인프라 운영에 강점을 가지고 있습니다. "
        "대용량 트래픽 처리와 마이크로서비스 아키텍처 설계 경험이 있으며, "
        "팀 협업과 코드 품질을 최우선으로 생각합니다. "
        "새로운 기술을 빠르게 습득하고 실무에 적용하는 능력을 갖추고 있습니다."
    )
    i_run = intro_para.add_run(intro_text)
    set_korean_font(i_run, size=11)

    # ── 경력사항 ───────────────────────────────────────────────────────────
    add_section_heading(doc, "경력사항")

    career_data = [
        ("회사명", "기간", "담당업무"),
        ("(주)테크스타트", "2022.03 ~ 현재", "백엔드 API 서버 개발 및 운영\nAWS 기반 클라우드 인프라 관리\nCI/CD 파이프라인 구축 (GitHub Actions)"),
        ("네오소프트 주식회사", "2020.07 ~ 2022.02", "사내 ERP 시스템 개발 및 유지보수\nMySQL 데이터베이스 설계 및 최적화\n레거시 시스템 리팩토링"),
        ("스마트로직스", "2019.01 ~ 2020.06", "물류 관리 시스템 백엔드 개발\nRESTful API 설계 및 문서화\n단위/통합 테스트 코드 작성"),
    ]

    career_para = doc.add_paragraph()
    career_para.paragraph_format.space_before = Pt(6)
    career_para.paragraph_format.space_after = Pt(0)

    table = doc.add_table(rows=len(career_data), cols=3)
    table.style = "Table Grid"

    # 열 너비
    col_widths = [Cm(4.5), Cm(3.8), Cm(8.0)]
    for i, width in enumerate(col_widths):
        for cell in table.columns[i].cells:
            cell.width = width

    # 헤더 행
    header_row = table.rows[0]
    header_texts = career_data[0]
    header_color = "1F497D"
    for j, text in enumerate(header_texts):
        cell = header_row.cells[j]
        set_cell_bg(cell, header_color)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.paragraphs[0].paragraph_format.space_before = Pt(4)
        cell.paragraphs[0].paragraph_format.space_after = Pt(4)
        run = cell.paragraphs[0].add_run(text)
        set_korean_font(run, size=11, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))

    # 데이터 행
    for i, row_data in enumerate(career_data[1:], start=1):
        row = table.rows[i]
        # 배경색 교대
        bg = "EEF3F9" if i % 2 == 1 else "FFFFFF"
        for j, text in enumerate(row_data):
            cell = row.cells[j]
            set_cell_bg(cell, bg)
            cell.paragraphs[0].paragraph_format.space_before = Pt(4)
            cell.paragraphs[0].paragraph_format.space_after = Pt(4)
            if j == 0:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            # 담당업무 칸: 줄바꿈 처리
            lines = text.split("\n")
            for k, line in enumerate(lines):
                if k == 0:
                    run = cell.paragraphs[0].add_run(line)
                else:
                    new_p = cell.add_paragraph(line)
                    new_p.paragraph_format.space_before = Pt(0)
                    new_p.paragraph_format.space_after = Pt(0)
                    run = new_p.runs[0] if new_p.runs else new_p.add_run(line)
                set_korean_font(run, size=10)

    # 표 다음 여백
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(4)

    # ── 학력 ───────────────────────────────────────────────────────────────
    add_section_heading(doc, "학력")

    education_items = [
        ("2015.03 ~ 2019.02", "한국대학교 컴퓨터공학과", "학사 졸업", "GPA 3.8 / 4.5"),
        ("2012.03 ~ 2015.02", "서울고등학교", "이과 졸업", ""),
    ]

    for edu in education_items:
        period, school, degree, note = edu
        edu_para = doc.add_paragraph()
        edu_para.paragraph_format.space_before = Pt(5)
        edu_para.paragraph_format.space_after = Pt(1)

        school_run = edu_para.add_run(f"{school}  ")
        set_korean_font(school_run, size=11, bold=True)

        degree_run = edu_para.add_run(f"({degree})")
        set_korean_font(degree_run, size=11)

        period_run = edu_para.add_run(f"    {period}")
        set_korean_font(period_run, size=10, color=RGBColor(0x77, 0x77, 0x77))

        if note:
            note_para = doc.add_paragraph()
            note_para.paragraph_format.space_before = Pt(0)
            note_para.paragraph_format.space_after = Pt(2)
            note_para.paragraph_format.left_indent = Cm(0.5)
            note_run = note_para.add_run(f"  {note}")
            set_korean_font(note_run, size=10, color=RGBColor(0x55, 0x55, 0x55))

    # ── 보유 기술 ──────────────────────────────────────────────────────────
    add_section_heading(doc, "보유 기술")

    skills = {
        "프로그래밍 언어": ["Python", "Java", "JavaScript", "SQL"],
        "프레임워크 / 라이브러리": ["FastAPI", "Spring Boot", "Django", "React"],
        "데이터베이스": ["PostgreSQL", "MySQL", "Redis", "MongoDB"],
        "클라우드 / 인프라": ["AWS (EC2, RDS, S3, Lambda)", "Docker", "Kubernetes", "Terraform"],
        "협업 / 도구": ["Git / GitHub", "Jira", "Confluence", "Slack"],
        "언어": ["한국어 (모국어)", "영어 (비즈니스 레벨)"],
    }

    for category, items in skills.items():
        skill_para = doc.add_paragraph()
        skill_para.paragraph_format.space_before = Pt(5)
        skill_para.paragraph_format.space_after = Pt(1)

        cat_run = skill_para.add_run(f"{category}: ")
        set_korean_font(cat_run, size=11, bold=True)

        items_run = skill_para.add_run("  /  ".join(items))
        set_korean_font(items_run, size=11)

    # ── 저장 ───────────────────────────────────────────────────────────────
    doc.save(OUTPUT_PATH)
    print(f"resume.docx 저장 완료: {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
