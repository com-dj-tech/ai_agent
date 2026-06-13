from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_PATH = "C:/Users/SBS/Desktop/agent02/docx_creator-workspace/iteration-1/eval-1/without_skill/outputs/employee_list.docx"

DARK_BLUE = RGBColor(0x1F, 0x39, 0x7A)  # #1F397A
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

HEADERS = ["이름", "부서", "직급", "입사일", "연락처"]

EMPLOYEES = [
    ["김민준", "개발팀", "선임연구원", "2019-03-15", "010-1234-5678"],
    ["이서연", "마케팅팀", "과장", "2020-07-01", "010-2345-6789"],
    ["박지훈", "인사팀", "대리", "2021-11-22", "010-3456-7890"],
    ["최유나", "재무팀", "차장", "2017-05-10", "010-4567-8901"],
    ["정현우", "영업팀", "사원", "2023-02-28", "010-5678-9012"],
]


def set_cell_background(cell, hex_color):
    """Set cell background color via XML shading."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_border(cell):
    """Add thin borders to a cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "AAAAAA")
        tcBorders.append(border)
    tcPr.append(tcBorders)


def main():
    doc = Document()

    # Page margins
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    # Title
    title = doc.add_heading("직원 명단", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.font.color.rgb = DARK_BLUE

    doc.add_paragraph()

    # Table: 6 rows (1 header + 5 data), 5 columns
    table = doc.add_table(rows=1 + len(EMPLOYEES), cols=len(HEADERS))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # Column widths
    col_widths = [Cm(3), Cm(3.5), Cm(3), Cm(3.5), Cm(4)]

    # --- Header row ---
    header_row = table.rows[0]
    for col_idx, (header_text, width) in enumerate(zip(HEADERS, col_widths)):
        cell = header_row.cells[col_idx]
        cell.width = width
        set_cell_background(cell, "1F397A")

        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(header_text)
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = WHITE

    # --- Data rows ---
    for row_idx, employee in enumerate(EMPLOYEES):
        row = table.rows[row_idx + 1]
        # Alternate row shading
        bg_color = "EEF2FB" if row_idx % 2 == 0 else "FFFFFF"

        for col_idx, value in enumerate(employee):
            cell = row.cells[col_idx]
            cell.width = col_widths[col_idx]
            set_cell_background(cell, bg_color)

            para = cell.paragraphs[0]
            # Center-align all columns
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(value)
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    doc.save(OUTPUT_PATH)
    print(f"Saved: {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
