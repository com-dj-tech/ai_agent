from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_PATH = r"C:/Users/SBS/Desktop/agent02/docx_creator-workspace/iteration-1/eval-1/with_skill/outputs/employee_list.docx"

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)

def set_cell_font(cell, bold=False, color_rgb=None, font_size=11):
    for para in cell.paragraphs:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in para.runs:
            run.bold = bold
            run.font.size = Pt(font_size)
            if color_rgb:
                run.font.color.rgb = color_rgb
            run.font.name = "맑은 고딕"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")

doc = Document()

# 페이지 설정 (A4)
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)

# 문서 제목
title = doc.add_heading("직원 명단", level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph("")  # 빈 줄

# 헤더 정의
headers = ["이름", "부서", "직급", "입사일", "연락처"]

# 샘플 데이터 5명
employees = [
    ["김민준", "개발팀",   "대리",   "2020-03-02", "010-1234-5678"],
    ["이서연", "인사팀",   "과장",   "2018-07-15", "010-2345-6789"],
    ["박준호", "마케팅팀", "사원",   "2023-01-10", "010-3456-7890"],
    ["최지아", "재무팀",   "차장",   "2015-09-20", "010-4567-8901"],
    ["정도윤", "기획팀",   "부장",   "2010-04-05", "010-5678-9012"],
]

# 표 생성 (헤더 1행 + 데이터 5행, 열 5개)
table = doc.add_table(rows=1 + len(employees), cols=len(headers))
table.style = "Table Grid"

# 열 너비 설정
col_widths = [Cm(3), Cm(3.5), Cm(3), Cm(3.5), Cm(4)]
for i, width in enumerate(col_widths):
    for cell in table.columns[i].cells:
        cell.width = width

# 헤더 행 설정 — 짙은 파란색 배경 + 흰색 굵은 글씨
DARK_BLUE = "003087"   # 짙은 파란색 hex
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

header_row = table.rows[0]
for j, h in enumerate(headers):
    cell = header_row.cells[j]
    cell.text = h
    set_cell_bg(cell, DARK_BLUE)
    set_cell_font(cell, bold=True, color_rgb=WHITE, font_size=11)

# 데이터 행 입력
for i, emp in enumerate(employees):
    row = table.rows[i + 1]
    for j, val in enumerate(emp):
        cell = row.cells[j]
        cell.text = val
        # 가운데 정렬 + 폰트
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.name = "맑은 고딕"
                run.font.size = Pt(10)
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")

doc.save(OUTPUT_PATH)
print(f"저장 완료: {OUTPUT_PATH}")
