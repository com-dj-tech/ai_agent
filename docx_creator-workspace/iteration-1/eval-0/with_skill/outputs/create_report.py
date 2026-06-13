from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_PATH = "C:/Users/SBS/Desktop/agent02/docx_creator-workspace/iteration-1/eval-0/with_skill/outputs/2025년_2분기_업무보고.docx"

doc = Document()

# ── 페이지 설정 (A4, 여백 2.5cm) ──────────────────────────────────────
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)

# ── 헤더: 부서명 ──────────────────────────────────────────────────────
header = section.header
header_para = header.paragraphs[0]
header_para.clear()
run_h = header_para.add_run("전략기획팀")
run_h.font.name = "맑은 고딕"
run_h._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
run_h.font.size = Pt(10)
run_h.font.bold = True
run_h.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

# ── 제목 ──────────────────────────────────────────────────────────────
title_para = doc.add_paragraph()
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_para.paragraph_format.space_before = Pt(12)
title_para.paragraph_format.space_after  = Pt(6)
run_title = title_para.add_run("2025년 2분기 업무보고")
run_title.bold = True
run_title.font.name = "맑은 고딕"
run_title._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
run_title.font.size = Pt(20)
run_title.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

# ── 작성일 ────────────────────────────────────────────────────────────
date_para = doc.add_paragraph()
date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_para.paragraph_format.space_after = Pt(18)
run_date = date_para.add_run("작성일: 2025-06-07")
run_date.font.name = "맑은 고딕"
run_date._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
run_date.font.size = Pt(11)
run_date.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

# ── 셀 배경색 유틸 ────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)

# ── 1. 요약 섹션 (파란색 배경 표) ────────────────────────────────────
section_heading = doc.add_heading("1. 요약", level=2)
section_heading.paragraph_format.space_before = Pt(12)
section_heading.paragraph_format.space_after  = Pt(6)

summary_data = [
    ("구분",          "내용"),
    ("보고 기간",     "2025년 4월 ~ 6월"),
    ("보고 부서",     "전략기획팀"),
    ("주요 과제",     "디지털 전환 추진 / 신규 서비스 기획 / 파트너십 확대"),
    ("달성률",        "목표 대비 92% 달성"),
    ("특이 사항",     "3분기 예산 재편성 필요"),
]

table = doc.add_table(rows=len(summary_data), cols=2)
table.style = "Table Grid"

HEADER_BG  = "1F497D"  # 진한 파란색 (헤더 행)
ROW_BG_ODD = "BDD7EE"  # 연한 파란색 (홀수 데이터 행)
ROW_BG_EVN = "DEEAF1"  # 더 연한 파란색 (짝수 데이터 행)

for r_idx, (col0, col1) in enumerate(summary_data):
    row = table.rows[r_idx]

    # 배경색 결정
    if r_idx == 0:
        bg0 = bg1 = HEADER_BG
        txt_color = RGBColor(0xFF, 0xFF, 0xFF)
        bold = True
    else:
        bg0 = bg1 = ROW_BG_ODD if r_idx % 2 == 1 else ROW_BG_EVN
        txt_color = RGBColor(0x1A, 0x1A, 0x1A)
        bold = False

    for c_idx, text in enumerate([col0, col1]):
        cell = row.cells[c_idx]
        set_cell_bg(cell, bg0 if c_idx == 0 else bg1)
        cell.paragraphs[0].clear()
        run = cell.paragraphs[0].add_run(text)
        run.bold = bold
        run.font.name = "맑은 고딕"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
        run.font.size = Pt(10)
        run.font.color.rgb = txt_color
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        cell.paragraphs[0].paragraph_format.space_before = Pt(3)
        cell.paragraphs[0].paragraph_format.space_after  = Pt(3)

# 열 너비
table.columns[0].width = Cm(4)
table.columns[1].width = Cm(11.5)

doc.add_paragraph()  # 간격

# ── 2. 주요 성과 (번호 목록) ─────────────────────────────────────────
doc.add_heading("2. 주요 성과", level=2)

achievements = [
    "디지털 전환 1단계 완료 — 핵심 업무 시스템 클라우드 이전 완료 (목표 대비 100%)",
    "신규 서비스 기획안 확정 — B2B SaaS 플랫폼 MVP 기획 완료, 개발팀 이관",
    "파트너십 3건 신규 체결 — A사, B사, C사와 전략적 제휴 MOU 서명",
    "내부 역량 강화 — 팀원 대상 데이터 분석 교육 2회 실시(참여율 95%)",
    "비용 절감 달성 — 운영비 전년 동기 대비 8.3% 절감",
]

for item in achievements:
    p = doc.add_paragraph(item, style="List Number")
    p.paragraph_format.space_after = Pt(4)
    for run in p.runs:
        run.font.name = "맑은 고딕"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
        run.font.size = Pt(10.5)

doc.add_paragraph()

# ── 3. 다음 계획 (글머리 기호 목록) ─────────────────────────────────
doc.add_heading("3. 다음 계획 (3분기)", level=2)

plans = [
    "디지털 전환 2단계 착수 — 외부 시스템 연동 및 데이터 파이프라인 구축",
    "신규 서비스 개발 착수 — MVP 개발 착수 및 베타 테스터 모집",
    "파트너십 협업 구체화 — 공동 마케팅 및 기술 협력 계획 수립",
    "팀 역량 강화 — AI 활용 업무 자동화 워크숍 개최",
    "3분기 예산 재편성 — 신규 과제 반영한 예산 조정안 제출 (7월 말 기한)",
]

for item in plans:
    p = doc.add_paragraph(item, style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    for run in p.runs:
        run.font.name = "맑은 고딕"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
        run.font.size = Pt(10.5)

# ── 저장 ──────────────────────────────────────────────────────────────
doc.save(OUTPUT_PATH)
print(f"문서 저장 완료: {OUTPUT_PATH}")
