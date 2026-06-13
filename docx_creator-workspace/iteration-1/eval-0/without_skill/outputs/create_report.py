"""
2025년 2분기 업무보고 DOCX 생성 스크립트
사용 라이브러리: python-docx
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy
import os

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))
OUTPUT_FILE = os.path.join(OUTPUT_DIR, "2025년_2분기_업무보고.docx")


def set_cell_background(cell, hex_color: str):
    """셀 배경색 설정"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_border(cell, **kwargs):
    """셀 테두리 설정 (top, bottom, left, right)"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), kwargs.get(edge, "single"))
        tag.set(qn("w:sz"), "4")
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), "4472C4")
        tcBorders.append(tag)
    tcPr.append(tcBorders)


def add_header(doc, department: str):
    """헤더에 부서명 추가"""
    section = doc.sections[0]
    header = section.header
    header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    header_para.clear()
    header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header_para.add_run(department)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)
    run.font.bold = True


def create_report():
    doc = Document()

    # ── A4 / 여백 2.5cm 설정 ──────────────────────────────────────────────
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    for attr in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(section, attr, Cm(2.5))

    # ── 기본 스타일 폰트 설정 (한글 지원) ─────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = "맑은 고딕"
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")

    # ── 헤더 설정 ─────────────────────────────────────────────────────────
    add_header(doc, "전략기획팀")

    # ── 제목 ──────────────────────────────────────────────────────────────
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.space_after = Pt(6)
    title_run = title_para.add_run("2025년 2분기 업무보고")
    title_run.font.size = Pt(22)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0x1F, 0x39, 0x64)
    title_run.font.name = "맑은 고딕"
    title_run.element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")

    # ── 작성일 ────────────────────────────────────────────────────────────
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para.paragraph_format.space_after = Pt(16)
    date_run = date_para.add_run("작성일: 2025-06-07")
    date_run.font.size = Pt(11)
    date_run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
    date_run.font.name = "맑은 고딕"
    date_run.element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")

    # ── 구분선 ────────────────────────────────────────────────────────────
    hr_para = doc.add_paragraph()
    hr_para.paragraph_format.space_after = Pt(12)
    pPr = hr_para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "4472C4")
    pBdr.append(bottom)
    pPr.append(pBdr)

    # ══════════════════════════════════════════════════════════════════════
    # 1. 요약 섹션 (파란색 배경 표)
    # ══════════════════════════════════════════════════════════════════════
    section_heading = doc.add_paragraph()
    section_heading.paragraph_format.space_before = Pt(4)
    section_heading.paragraph_format.space_after = Pt(8)
    sh_run = section_heading.add_run("1. 업무 요약")
    sh_run.font.size = Pt(14)
    sh_run.font.bold = True
    sh_run.font.color.rgb = RGBColor(0x1F, 0x39, 0x64)
    sh_run.font.name = "맑은 고딕"
    sh_run.element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")

    summary_data = [
        ("구분", "내용"),
        ("보고 기간", "2025년 4월 1일 ~ 2025년 6월 30일"),
        ("담당 부서", "전략기획팀"),
        ("보고 목적", "2분기 주요 업무 실적 및 향후 계획 보고"),
        ("전반적 달성률", "목표 대비 92% 달성"),
    ]

    table = doc.add_table(rows=len(summary_data), cols=2)
    table.style = "Table Grid"
    table.autofit = False

    col_widths = [Cm(4.0), Cm(12.0)]
    for row_idx, (key, value) in enumerate(summary_data):
        row = table.rows[row_idx]
        row.height = Cm(0.9)

        # 헤더 행
        if row_idx == 0:
            for col_idx, text in enumerate((key, value)):
                cell = row.cells[col_idx]
                cell.width = col_widths[col_idx]
                set_cell_background(cell, "1F3964")
                set_cell_border(cell)
                para = cell.paragraphs[0]
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = para.add_run(text)
                run.font.bold = True
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.name = "맑은 고딕"
                run.element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
        else:
            # 항목 셀 (짙은 파랑)
            key_cell = row.cells[0]
            key_cell.width = col_widths[0]
            set_cell_background(key_cell, "BDD7EE")
            set_cell_border(key_cell)
            key_para = key_cell.paragraphs[0]
            key_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            key_run = key_para.add_run(key)
            key_run.font.bold = True
            key_run.font.size = Pt(10)
            key_run.font.color.rgb = RGBColor(0x1F, 0x39, 0x64)
            key_run.font.name = "맑은 고딕"
            key_run.element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")

            # 값 셀 (연한 파랑)
            val_cell = row.cells[1]
            val_cell.width = col_widths[1]
            set_cell_background(val_cell, "DEEAF1")
            set_cell_border(val_cell)
            val_para = val_cell.paragraphs[0]
            val_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            val_run = val_para.add_run(value)
            val_run.font.size = Pt(10)
            val_run.font.name = "맑은 고딕"
            val_run.element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")

    doc.add_paragraph()  # 간격

    # ══════════════════════════════════════════════════════════════════════
    # 2. 주요 성과 (번호 목록)
    # ══════════════════════════════════════════════════════════════════════
    perf_heading = doc.add_paragraph()
    perf_heading.paragraph_format.space_before = Pt(4)
    perf_heading.paragraph_format.space_after = Pt(8)
    ph_run = perf_heading.add_run("2. 주요 성과")
    ph_run.font.size = Pt(14)
    ph_run.font.bold = True
    ph_run.font.color.rgb = RGBColor(0x1F, 0x39, 0x64)
    ph_run.font.name = "맑은 고딕"
    ph_run.element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")

    achievements = [
        "신규 사업 기획안 3건 수립 및 경영진 승인 완료 (달성률 150%)",
        "전사 디지털 전환(DX) 로드맵 초안 작성 및 유관 부서 협의 완료",
        "2분기 전략 KPI 92% 달성 — 전 분기 대비 7%p 향상",
        "외부 파트너사 MOU 체결 2건 (A사, B사) 및 협력 과제 착수",
        "부서 내 업무 자동화 도구 도입으로 반복 업무 처리 시간 35% 단축",
        "임직원 전략 교육 프로그램 운영 (참가자 48명, 만족도 4.6/5.0)",
    ]

    for i, item in enumerate(achievements, 1):
        para = doc.add_paragraph(style="List Number")
        para.paragraph_format.space_after = Pt(4)
        para.paragraph_format.left_indent = Cm(0.5)
        run = para.add_run(f"{item}")
        run.font.size = Pt(10.5)
        run.font.name = "맑은 고딕"
        run.element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")

    doc.add_paragraph()  # 간격

    # ══════════════════════════════════════════════════════════════════════
    # 3. 다음 계획 (글머리 기호 목록)
    # ══════════════════════════════════════════════════════════════════════
    plan_heading = doc.add_paragraph()
    plan_heading.paragraph_format.space_before = Pt(4)
    plan_heading.paragraph_format.space_after = Pt(8)
    plh_run = plan_heading.add_run("3. 다음 계획 (3분기)")
    plh_run.font.size = Pt(14)
    plh_run.font.bold = True
    plh_run.font.color.rgb = RGBColor(0x1F, 0x39, 0x64)
    plh_run.font.name = "맑은 고딕"
    plh_run.element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")

    next_plans = [
        "3분기 전략 목표 수립 및 부서별 KPI 배분 (7월 첫째 주)",
        "신규 사업 파일럿 프로젝트 착수 — A사 공동 과제 실행 개시",
        "전사 DX 로드맵 최종 확정 및 이사회 보고",
        "하반기 예산 재배분 검토 및 CFO 협의",
        "전략 교육 2차 과정 기획 — 리더십 역량 강화 프로그램 추가",
        "글로벌 시장 조사 보고서 작성 및 해외 진출 가능성 검토",
        "부서 내 업무 프로세스 표준화 매뉴얼 완성 (8월 말 목표)",
    ]

    for item in next_plans:
        para = doc.add_paragraph(style="List Bullet")
        para.paragraph_format.space_after = Pt(4)
        para.paragraph_format.left_indent = Cm(0.5)
        run = para.add_run(item)
        run.font.size = Pt(10.5)
        run.font.name = "맑은 고딕"
        run.element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")

    # ── 하단 서명란 ───────────────────────────────────────────────────────
    doc.add_paragraph()
    sig_para = doc.add_paragraph()
    sig_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    sig_para.paragraph_format.space_before = Pt(24)
    sig_run = sig_para.add_run("전략기획팀장  ___________________  (인)")
    sig_run.font.size = Pt(10.5)
    sig_run.font.name = "맑은 고딕"
    sig_run.element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")

    # ── 저장 ──────────────────────────────────────────────────────────────
    doc.save(OUTPUT_FILE)
    print(f"[완료] 파일 저장: {OUTPUT_FILE}")


if __name__ == "__main__":
    create_report()
