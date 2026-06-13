from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)

def set_korean_font(run, size=None, bold=False, color=None):
    run.font.name = "맑은 고딕"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    if size:
        run.font.size = Pt(size)
    if bold:
        run.bold = True
    if color:
        run.font.color.rgb = color

doc = Document()

# 페이지 설정 (A4)
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin   = Cm(3)
section.right_margin  = Cm(3)

# 헤더
header_para = section.header.paragraphs[0]
header_para.text = "위키백과 요약 보고서"
header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
for run in header_para.runs:
    run.font.name = "맑은 고딕"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

# 푸터 페이지 번호
footer_para = section.footer.paragraphs[0]
footer_para.clear()
run = footer_para.add_run()
fldChar1 = OxmlElement("w:fldChar")
fldChar1.set(qn("w:fldCharType"), "begin")
instrText = OxmlElement("w:instrText")
instrText.text = "PAGE"
fldChar2 = OxmlElement("w:fldChar")
fldChar2.set(qn("w:fldCharType"), "end")
run._r.extend([fldChar1, instrText, fldChar2])
footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ─── 표지 영역 ───────────────────────────────────────────
title_para = doc.add_paragraph()
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_para.paragraph_format.space_before = Pt(40)
title_para.paragraph_format.space_after = Pt(6)
run = title_para.add_run("인공지능(AI) 위키백과 요약 보고서")
set_korean_font(run, size=22, bold=True, color=RGBColor(0x1F, 0x38, 0x64))

sub_para = doc.add_paragraph()
sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_para.paragraph_format.space_after = Pt(4)
run = sub_para.add_run("Artificial Intelligence — Wikipedia 한국어판 요약")
set_korean_font(run, size=12, color=RGBColor(0x44, 0x72, 0xC4))

# 구분선
hr = doc.add_paragraph()
hr.paragraph_format.space_before = Pt(6)
hr.paragraph_format.space_after = Pt(6)
hr_run = hr.add_run("─" * 55)
hr_run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

# 메타 정보 표
meta_table = doc.add_table(rows=3, cols=2)
meta_table.style = "Table Grid"
meta_data = [
    ("출처", "위키백과 한국어판 (ko.wikipedia.org/wiki/인공지능)"),
    ("작성일", "2026년 6월 13일"),
    ("마지막 편집", "2026년 5월 13일"),
]
for i, (key, val) in enumerate(meta_data):
    key_cell = meta_table.cell(i, 0)
    val_cell = meta_table.cell(i, 1)
    set_cell_bg(key_cell, "1F3864")
    key_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    kr = key_cell.paragraphs[0].add_run(key)
    set_korean_font(kr, size=10, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
    vr = val_cell.paragraphs[0].add_run(val)
    set_korean_font(vr, size=10)

meta_table.columns[0].width = Cm(3.5)
meta_table.columns[1].width = Cm(13)

doc.add_paragraph()

# ─── 1. 개요 ─────────────────────────────────────────────
doc.add_heading("1. 개요", level=2)
p = doc.add_paragraph()
r = p.add_run(
    "인공지능(人工智能, Artificial Intelligence, AI)은 인간의 학습능력, 추론능력, 지각능력을 인공적으로 구현하려는 "
    "컴퓨터 과학의 세부 분야이다. 정보공학 분야의 핵심 인프라 기술이며, 인간의 지능을 기계에 인공적으로 시연·구현한 "
    "컴퓨터 시스템이다. 인간을 포함한 동물이 갖고 있는 자연 지능(natural intelligence)과는 구별된다."
)
set_korean_font(r, size=11)
p.paragraph_format.line_spacing = Pt(20)

# ─── 2. 강인공지능 vs 약인공지능 ──────────────────────────
doc.add_heading("2. 강인공지능 vs 약인공지능", level=2)

items_ai = [
    ("약인공지능 (Weak AI)",
     "특정 문제 해결에 특화된 도구형 AI. 사진 인식, 음성 인식 등 실용적·현실적 목표에 집중. "
     "기존에 인간은 쉽게 해결하나 컴퓨터로 처리하기 어려웠던 문제를 수행하는 데 중점을 둔다."),
    ("강인공지능 (Strong AI / AGI)",
     "인간처럼 실제로 사고하여 문제를 해결하는 '일반 지능' 구현 시도. 추론·문제 해결·지식 표현·계획 수립·학습 등 "
     "통합 능력이 필요하며, 새로운 환경에서도 스스로 학습·적응하는 일반화 능력이 핵심 요소이다."),
]
for label, desc in items_ai:
    p = doc.add_paragraph(style="List Bullet")
    r_label = p.add_run(label + ": ")
    set_korean_font(r_label, size=11, bold=True)
    r_desc = p.add_run(desc)
    set_korean_font(r_desc, size=11)

# ─── 3. 역사 연표 ─────────────────────────────────────────
doc.add_heading("3. 역사 연표", level=2)

history_data = [
    ("시기", "주요 사건"),
    ("1943~1956", "AI 이론 태동. 앨런 튜링의 튜링 테스트 제안(1950), 다트머스 컨퍼런스(1956)에서 AI 학문 분야 공식 출범"),
    ("1956~1974", "황금기. 대수·기하 문제 풀기, 자연어 처리(ELIZA 채팅 프로그램), 로보틱스 등 급속한 발전"),
    ("1974~1980", "첫 번째 암흑기. 복잡한 문제 해결 실패, 연구 한계 노출, 정부·기업 자금 지원 급감"),
    ("1980~1987", "AI 붐. 전문가 시스템 상용화, 일본 제5세대 컴퓨터 프로젝트, DARPA 투자 재개"),
    ("1987~1993", "두 번째 암흑기. 전문가 시스템 유지비 급증, 시장 붕괴, 연결주의(신경망) 재등장"),
    ("1993~현재", "AI 부활. 머신러닝·딥러닝 발전, 빅데이터 활용, GPU 가속, 생성형 AI(LLM) 등장"),
]

table = doc.add_table(rows=len(history_data), cols=2)
table.style = "Table Grid"
for i, (col1, col2) in enumerate(history_data):
    c1 = table.cell(i, 0)
    c2 = table.cell(i, 1)
    if i == 0:
        set_cell_bg(c1, "1F3864")
        set_cell_bg(c2, "1F3864")
        c1.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        c2.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = c1.paragraphs[0].add_run(col1)
        r2 = c2.paragraphs[0].add_run(col2)
        set_korean_font(r1, size=10, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        set_korean_font(r2, size=10, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
    else:
        if i % 2 == 0:
            set_cell_bg(c1, "EBF3FB")
            set_cell_bg(c2, "EBF3FB")
        c1.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = c1.paragraphs[0].add_run(col1)
        r2 = c2.paragraphs[0].add_run(col2)
        set_korean_font(r1, size=10, bold=True)
        set_korean_font(r2, size=10)

table.columns[0].width = Cm(3.5)
table.columns[1].width = Cm(13)

# ─── 4. 한계와 문제점 ─────────────────────────────────────
doc.add_heading("4. 인공지능의 한계와 문제점", level=2)

limits = [
    ("거짓정보 전달 및 가짜뉴스", "AI가 생성한 허위 정보 확산 위험. 딥페이크, 자동 생성 기사 등 악용 가능성"),
    ("인간의 통제력 약화", "AI 의사결정 증가로 인한 인간 판단력 저하 및 책임 소재 불명확"),
    ("일자리 감소", "자동화로 인한 노동 시장 구조 변화. 반복적 업무의 대체 가속화"),
    ("윤리 문제", "편향된 학습 데이터로 인한 차별적 결과 도출 가능성"),
    ("개인정보 유출", "방대한 데이터 수집·처리 과정에서의 프라이버시 침해"),
    ("사고력 저하", "AI 의존도 증가로 인한 인간 사고 능력 감퇴 우려"),
]
for label, desc in limits:
    p = doc.add_paragraph(style="List Bullet")
    r_label = p.add_run(label + ": ")
    set_korean_font(r_label, size=11, bold=True)
    r_desc = p.add_run(desc)
    set_korean_font(r_desc, size=11)

# ─── 5. 응용 분야 ─────────────────────────────────────────
doc.add_heading("5. 실용적 응용 분야", level=2)

apps = [
    "자연어 처리 — 번역, 챗봇, 텍스트 생성 (GPT, BERT 등)",
    "컴퓨터 비전 — 이미지 인식, 객체 감지, 자율주행",
    "의료 — 진단 보조, 신약 개발, 의료 영상 분석",
    "게임 AI — 체스(Deep Blue), 바둑(AlphaGo), 전략 시뮬레이션",
    "로보틱스 및 자동화 — 제조·물류 자동화, 서비스 로봇",
]
for item in apps:
    p = doc.add_paragraph(item, style="List Bullet")
    for run in p.runs:
        set_korean_font(run, size=11)

# ─── 6. 유명 AI 및 연구자 ─────────────────────────────────
doc.add_heading("6. 유명 인공지능 및 연구자", level=2)

notable = [
    ("지능적 기계",
     "Deep Blue(IBM, 체스), AlphaGo(구글 딥마인드, 바둑), GPT 시리즈(OpenAI, 자연어 생성), ELIZA(초기 챗봇)"),
    ("주요 연구자",
     "존 매카시(AI 명명·다트머스 회의), 마빈 민스키(신경망·SNARC), 앨런 튜링(튜링 테스트), "
     "앨런 뉴얼·허버트 사이먼(General Problem Solver), 제프리 힌턴(딥러닝)"),
]
for label, desc in notable:
    p = doc.add_paragraph(style="List Bullet")
    r_label = p.add_run(label + ": ")
    set_korean_font(r_label, size=11, bold=True)
    r_desc = p.add_run(desc)
    set_korean_font(r_desc, size=11)

# ─── 7. 미래 전망 ─────────────────────────────────────────
doc.add_heading("7. 미래 전망", level=2)

future_items = [
    ("초지능 (Superintelligence)", "인간 지능을 초월하는 AI 출현 가능성 논의. 닉 보스트롬 등 연구자들이 위험성 경고"),
    ("AGI 개발", "범용 인공지능(AGI) 실현을 위한 글로벌 경쟁 심화. OpenAI, DeepMind, Meta 등 대규모 투자"),
    ("규제 및 안전", "통제 불능 AI, 악용 가능성에 대한 글로벌 규제 논의 증가. EU AI Act 등 법제화 진행"),
    ("안전한 AI 개발", "윤리적·안전한 AI 개발이 핵심 과제로 부상. 설명 가능한 AI(XAI) 연구 확대"),
]
for label, desc in future_items:
    p = doc.add_paragraph(style="List Bullet")
    r_label = p.add_run(label + ": ")
    set_korean_font(r_label, size=11, bold=True)
    r_desc = p.add_run(desc)
    set_korean_font(r_desc, size=11)

# 마무리 구분선
doc.add_paragraph()
hr2 = doc.add_paragraph()
hr2.alignment = WD_ALIGN_PARAGRAPH.CENTER
hr2_run = hr2.add_run("─" * 55)
hr2_run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

note = doc.add_paragraph()
note.alignment = WD_ALIGN_PARAGRAPH.CENTER
nr = note.add_run("본 보고서는 위키백과 한국어판(ko.wikipedia.org)의 '인공지능' 문서를 요약·정리한 것입니다.")
set_korean_font(nr, size=9, color=RGBColor(0x88, 0x88, 0x88))

# 저장
output_path = r"C:/Users/SBS/Desktop/agent02/PW/인공지능_위키백과_요약보고서.docx"
doc.save(output_path)
print(f"저장 완료: {output_path}")
