from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)

def set_run_font(run, name="맑은 고딕", size=11, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

doc = Document()

# 페이지 설정
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin   = Cm(3)
section.right_margin  = Cm(3)

# 제목
title = doc.add_heading("네이버 기업 분석 보고서", level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.name = "맑은 고딕"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    run.font.color.rgb = RGBColor(0x03, 0xC7, 0x5A)

sub = doc.add_paragraph("작성일: 2026년 6월 | 출처: NAVER Corp. 공식 발표, Wikipedia, PitchBook")
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in sub.runs:
    set_run_font(run, size=10, color=(120, 120, 120))

doc.add_paragraph()

# ── 1. 기업 개요 ──────────────────────────────────────────
doc.add_heading("1. 기업 개요", level=2)

overview_data = [
    ("정식 명칭", "NAVER Corporation"),
    ("설립연도", "1999년"),
    ("본사 소재지", "경기도 성남시 분당구 정자동 (그린팩토리)"),
    ("대표이사 (CEO)", "최수연"),
    ("상장 거래소", "한국거래소 (KRX) KOSPI · 종목코드: 035420"),
    ("시가총액", "약 27.7조 원 (2026년 3월 기준)"),
    ("연간 매출액", "12조 350억 원 (2025년 연간)"),
    ("사업 영역", "검색 포털, 커머스, 핀테크, 콘텐츠, 클라우드, AI"),
]

table = doc.add_table(rows=len(overview_data), cols=2)
table.style = "Table Grid"
for i, (key, val) in enumerate(overview_data):
    row = table.rows[i]
    row.cells[0].text = key
    row.cells[1].text = val
    set_cell_bg(row.cells[0], "E8F5E9")
    for cell in row.cells:
        for para in cell.paragraphs:
            for run in para.runs:
                set_run_font(run, size=10)
    table.columns[0].width = Cm(5)
    table.columns[1].width = Cm(10)

doc.add_paragraph()

# ── 2. 연혁 ──────────────────────────────────────────────
doc.add_heading("2. 주요 연혁", level=2)

milestones = [
    ("1999", "NAVER Corporation 창립 (이해진 창업자)"),
    ("2000", "'통합검색' 서비스 출시 — 단일 페이지에서 다양한 결과 유형 제공"),
    ("2002", "KOSDAQ 상장 / 지식iN 서비스 출시"),
    ("2004", "웹툰 서비스 출시 (디지털 만화 플랫폼 선도)"),
    ("2011", "모바일 메신저 LINE 일본 출시"),
    ("2013", "한게임과 분리, NAVER Corporation으로 사명 변경"),
    ("2017", "SNOW, Naver Labs, Naver Webtoon 자회사 설립"),
    ("2024", "Webtoon Entertainment, 나스닥 IPO (기업가치 약 2.9조 원)"),
    ("2025", "CEO 최수연 연임 확정 / 연간 매출 12조 원 돌파"),
    ("2026", "사업 부문 재편: 네이버플랫폼·금융플랫폼·글로벌챌린지 3분류 체계"),
]

table2 = doc.add_table(rows=len(milestones) + 1, cols=2)
table2.style = "Table Grid"
headers = ["연도", "주요 사항"]
for j, h in enumerate(headers):
    cell = table2.rows[0].cells[j]
    cell.text = h
    set_cell_bg(cell, "03C75A")
    for para in cell.paragraphs:
        for run in para.runs:
            set_run_font(run, size=10, bold=True, color=(255, 255, 255))

for i, (year, event) in enumerate(milestones):
    row = table2.rows[i + 1]
    row.cells[0].text = year
    row.cells[1].text = event
    if i % 2 == 0:
        set_cell_bg(row.cells[0], "F1F8F4")
        set_cell_bg(row.cells[1], "F1F8F4")
    for cell in row.cells:
        for para in cell.paragraphs:
            for run in para.runs:
                set_run_font(run, size=10)

doc.add_paragraph()

# ── 3. 사업 부문 ──────────────────────────────────────────
doc.add_heading("3. 사업 부문", level=2)

segments = [
    ("검색 플랫폼", "2025 Q4: 1조 596억 원", "네이버 검색, 디스플레이 광고, Whale 브라우저, CLOVA AI 검색"),
    ("커머스", "2025 Q4: 1조 540억 원", "네이버쇼핑, 스마트스토어, KREAM, 브랜드스토어"),
    ("핀테크", "2025 Q4: 4,531억 원", "네이버페이, 네이버파이낸셜, 간편결제·대출·보험"),
    ("콘텐츠", "2025 Q4: 4,567억 원", "웹툰, CHZZK(라이브스트리밍), 오디오클립, 시리즈"),
    ("클라우드·AI", "2025 Q4: 1,718억 원", "네이버 클라우드, HyperCLOVA X, 데이터센터"),
]

table3 = doc.add_table(rows=len(segments) + 1, cols=3)
table3.style = "Table Grid"
for j, h in enumerate(["사업 부문", "2025 Q4 매출", "주요 서비스"]):
    cell = table3.rows[0].cells[j]
    cell.text = h
    set_cell_bg(cell, "1A5276")
    for para in cell.paragraphs:
        for run in para.runs:
            set_run_font(run, size=10, bold=True, color=(255, 255, 255))

for i, (seg, rev, desc) in enumerate(segments):
    row = table3.rows[i + 1]
    row.cells[0].text = seg
    row.cells[1].text = rev
    row.cells[2].text = desc
    bg = "EBF5FB" if i % 2 == 0 else "FDFEFE"
    for cell in row.cells:
        set_cell_bg(cell, bg)
        for para in cell.paragraphs:
            for run in para.runs:
                set_run_font(run, size=10)

doc.add_paragraph()

# ── 4. 재무 실적 ──────────────────────────────────────────
doc.add_heading("4. 재무 실적", level=2)

fin_data = [
    ("구분", "2025 연간", "2026 Q1", "전년 동기 대비"),
    ("연결 매출액", "12조 350억 원", "3조 2,411억 원", "+16.3%"),
    ("영업이익", "2조 2,081억 원", "5,418억 원", "+7.2%"),
    ("시가총액", "—", "약 27.7조 원", "—"),
    ("주당 순이익", "—", "공시 예정", "—"),
]

table4 = doc.add_table(rows=len(fin_data), cols=4)
table4.style = "Table Grid"
for i, row_data in enumerate(fin_data):
    for j, val in enumerate(row_data):
        cell = table4.rows[i].cells[j]
        cell.text = val
        if i == 0:
            set_cell_bg(cell, "2C3E50")
            for para in cell.paragraphs:
                for run in para.runs:
                    set_run_font(run, size=10, bold=True, color=(255, 255, 255))
        else:
            for para in cell.paragraphs:
                for run in para.runs:
                    set_run_font(run, size=10)

doc.add_paragraph()
note = doc.add_paragraph("※ 2026년 Q1부터 사업 부문 분류 체계가 네이버플랫폼 / 금융플랫폼 / 글로벌챌린지 3분류로 재편됨.")
for run in note.runs:
    set_run_font(run, size=9, color=(100, 100, 100))

doc.add_paragraph()

# ── 5. 주요 서비스 ────────────────────────────────────────
doc.add_heading("5. 주요 서비스 및 브랜드", level=2)

services = [
    ("네이버 검색", "국내 시장점유율 1위 검색 포털. 통합검색·실시간 이슈·AI 검색(CLOVA X) 제공."),
    ("웹툰 (Webtoon)", "2004년 출시. 세계 최대 디지털 만화 플랫폼. 2024년 나스닥 상장."),
    ("LINE", "일본·동남아 주요 메신저. LY Corporation(LINE+Yahoo Japan 합병법인) 지분 32.5% 보유."),
    ("CLOVA AI / HyperCLOVA X", "자체 개발 대형 언어 모델(LLM). 검색, 챗봇, B2B 솔루션에 적용."),
    ("네이버페이", "간편결제·금융 서비스. 누적 가입자 수천만 명 규모의 핀테크 플랫폼."),
    ("네이버 클라우드", "B2B 클라우드·데이터센터 서비스. AI 인프라 및 엔터프라이즈 솔루션 제공."),
    ("CHZZK", "라이브 스트리밍 플랫폼. 게임·엔터테인먼트 중심으로 급성장."),
    ("KREAM", "한정판 스니커즈·패션 리세일 C2C 플랫폼."),
    ("SNOW", "카메라·AR 필터 앱. 글로벌 사용자 기반 보유."),
]

for svc, desc in services:
    para = doc.add_paragraph(style="List Bullet")
    run_title = para.add_run(f"{svc}: ")
    set_run_font(run_title, size=10, bold=True)
    run_desc = para.add_run(desc)
    set_run_font(run_desc, size=10)

doc.add_paragraph()

# ── 6. 경영진 ─────────────────────────────────────────────
doc.add_heading("6. 주요 경영진", level=2)

mgmt = [
    ("이해진", "글로벌 투자 책임자 (GIO) / 창업자"),
    ("최수연", "대표이사 (CEO) — 2025년 2월 연임 확정"),
    ("변대규", "이사회 의장 (Chairman)"),
    ("준구 김 (Junkoo Kim)", "Webtoon Entertainment CEO"),
]

table5 = doc.add_table(rows=len(mgmt) + 1, cols=2)
table5.style = "Table Grid"
for j, h in enumerate(["성명", "직책"]):
    cell = table5.rows[0].cells[j]
    cell.text = h
    set_cell_bg(cell, "03C75A")
    for para in cell.paragraphs:
        for run in para.runs:
            set_run_font(run, size=10, bold=True, color=(255, 255, 255))

for i, (name, role) in enumerate(mgmt):
    row = table5.rows[i + 1]
    row.cells[0].text = name
    row.cells[1].text = role
    for cell in row.cells:
        for para in cell.paragraphs:
            for run in para.runs:
                set_run_font(run, size=10)

doc.add_paragraph()

# 푸터
section = doc.sections[0]
footer = section.footer
fp = footer.paragraphs[0]
fp.text = "© 2026 네이버 기업 분석 보고서 | 본 문서는 공개 정보를 기반으로 작성되었습니다."
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in fp.runs:
    set_run_font(run, size=9, color=(150, 150, 150))

doc.save("C:/Users/SBS/Desktop/agent02/web_maker/naver_research.docx")
print("저장 완료: web_maker/naver_research.docx")
