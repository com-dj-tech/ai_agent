from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DOC_PATH = "C:/Users/SBS/Desktop/agent02/web_maker/naver_research.docx"

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)

def set_run_font(run, name="맑은 고딕", size=11, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

doc = Document(DOC_PATH)

# 기존 푸터 제거 후 페이지 구분
doc.add_page_break()

# ══════════════════════════════════════════════════
# 7. 디자인 컨셉
# ══════════════════════════════════════════════════
doc.add_heading("7. 디자인 컨셉 및 브랜드 아이덴티티", level=2)

intro = doc.add_paragraph(
    "네이버의 디자인 철학은 '일관성(Consistency)', '접근성(Accessibility)', '기술과 인간의 조화(Human-Tech Harmony)'를 "
    "핵심 축으로 삼습니다. 브랜드 컬러·타이포그래피·공간 디자인·아이콘 시스템에 이르기까지 모든 접점에서 "
    "통일된 시각 언어를 구현합니다."
)
for run in intro.runs:
    set_run_font(run, size=10)

doc.add_paragraph()

# 7-1. 브랜드 컬러
doc.add_heading("7-1. 브랜드 컬러 — 네이버 그린", level=3)

color_data = [
    ("공식 컬러명", "NAVER Green"),
    ("HEX 코드", "#03C75A"),
    ("RGB 값", "R 3 · G 199 · B 90"),
    ("상징 의미", "신뢰, 친근함, 탐험, 혁신"),
    ("적용 방침", "모든 사용자 접점에 일관 적용 — 앱 아이콘, 로고, CTA 버튼, 사옥 외관"),
]

table_color = doc.add_table(rows=len(color_data), cols=2)
table_color.style = "Table Grid"
for i, (k, v) in enumerate(color_data):
    row = table_color.rows[i]
    row.cells[0].text = k
    row.cells[1].text = v
    set_cell_bg(row.cells[0], "E8F5E9")
    for cell in row.cells:
        for para in cell.paragraphs:
            for run in para.runs:
                set_run_font(run, size=10)
    table_color.columns[0].width = Cm(5)
    table_color.columns[1].width = Cm(10)

doc.add_paragraph()

note_color = doc.add_paragraph(
    "※ 로고타입은 그래픽 요소 없이 고딕 계열 서체만 사용하며, 굵고 직선적인 획과 기하학적 형태로 "
    "안정감·가독성·단순미를 표현합니다. (출처: navercorp.com/en/company/brandGuide)"
)
for run in note_color.runs:
    set_run_font(run, size=9, color=(100, 100, 100))

doc.add_paragraph()

# 7-2. 타이포그래피
doc.add_heading("7-2. 타이포그래피", level=3)

typo_items = [
    ("나눔폰트 (Nanum Font)",
     "2008년 한글날 '한글한글 아름답게' 캠페인으로 무료 배포 시작. "
     "나눔고딕·나눔명조·나눔바른고딕 등으로 구성. 한국적 조형미와 현대적 미감을 결합한 무료 오픈소스 폰트."),
    ("나눔고딕 코딩",
     "개발자 커뮤니티를 위한 고정폭(monospace) 한글 코딩 폰트. GitHub 오픈소스 공개."),
    ("적용 원칙",
     "가독성 최우선 / 동아시아(한글) 폰트와 서구 폰트의 일관된 혼용 / "
     "모바일·웹·인쇄 환경 모두 대응하는 다중 굵기(Weight) 체계 운영."),
]

for title, desc in typo_items:
    para = doc.add_paragraph(style="List Bullet")
    run_t = para.add_run(f"{title}: ")
    set_run_font(run_t, size=10, bold=True)
    run_d = para.add_run(desc)
    set_run_font(run_d, size=10)

doc.add_paragraph()

# 7-3. 아이콘 시스템
doc.add_heading("7-3. 아이콘 시스템 — 189개 통합 리뉴얼", level=3)

icon_data = [
    ("리뉴얼 규모", "전체 서비스 189개 아이콘"),
    ("소요 기간", "468일"),
    ("착수 시점", "2021년 트렌드 변화에 따른 전면 개편 결정"),
    ("핵심 원칙", "배경과 메타포 분리 — 흰 배경에 상징 오브젝트만 배치, 어떤 환경에서도 활용 가능"),
    ("시각 특성", "2.5D 공간감 / 명확한 명암 대비 / 색상과 그림자의 구조적 가독성 향상"),
    ("컬러 변화", "NAVER 그린 미세 조정 — 이전 아이콘 대비 더 선명하고 구조적인 색상 표현"),
]

table_icon = doc.add_table(rows=len(icon_data), cols=2)
table_icon.style = "Table Grid"
for i, (k, v) in enumerate(icon_data):
    row = table_icon.rows[i]
    row.cells[0].text = k
    row.cells[1].text = v
    set_cell_bg(row.cells[0], "EBF5FB")
    for cell in row.cells:
        for para in cell.paragraphs:
            for run in para.runs:
                set_run_font(run, size=10)
    table_icon.columns[0].width = Cm(4.5)
    table_icon.columns[1].width = Cm(10.5)

doc.add_paragraph()

# 7-4. 공간 디자인 — 1784
doc.add_heading("7-4. 공간 디자인 — 제2사옥 '1784'", level=3)

space_intro = doc.add_paragraph(
    "2022년 4월 개관한 네이버 제2사옥 '1784'는 iF 디자인 어워드를 수상한 세계 최초 로봇 친화형 건물입니다. "
    "'기술 융합(Technological Convergence)'을 핵심 컨셉으로, 건축·인테리어·UX가 유기적으로 통합된 "
    "리빙랩(Living Lab)으로 설계되었습니다."
)
for run in space_intro.runs:
    set_run_font(run, size=10)

doc.add_paragraph()

space_data = [
    ("설계사", "삼성물산 삼우종합건축사사무소 (Samoo)"),
    ("규모", "지하 8층 · 지상 28층 / 연면적 168,179㎡"),
    ("핵심 컨셉", "기술 융합 (Technological Convergence)"),
    ("소재·마감", "철골·콘크리트 중심의 미래지향적 질감 + 목재 바닥·녹지로 온기 부여"),
    ("로봇 인프라",
     "루키(Rookie) 로봇 100대 운영 / 로봇 전용 엘리베이터 설치 / "
     "건물 자체 5G 네트워크 구축 (정부 특별 허가) / 디지털 트윈으로 6인치 정밀도 위치 추적"),
    ("스마트 오피스",
     "5,000명 임직원 + 100대 로봇 협업 공간 / 카페·배달·물류를 로봇이 자율 수행"),
    ("디자인 수상", "iF Design Award 수상 (NAVER 1784 부문)"),
    ("미디어 시스템",
     "건물 내 대형 미디어 파사드 및 디지털 사이니지 통합 설계 (NAVER 1784 Media System Design)"),
]

table_space = doc.add_table(rows=len(space_data), cols=2)
table_space.style = "Table Grid"
for i, (k, v) in enumerate(space_data):
    row = table_space.rows[i]
    row.cells[0].text = k
    row.cells[1].text = v
    if i % 2 == 0:
        set_cell_bg(row.cells[0], "F4ECF7")
        set_cell_bg(row.cells[1], "F4ECF7")
    for cell in row.cells:
        for para in cell.paragraphs:
            for run in para.runs:
                set_run_font(run, size=10)
    table_space.columns[0].width = Cm(4.5)
    table_space.columns[1].width = Cm(10.5)

doc.add_paragraph()

# 7-5. 서비스별 디자인 이니셔티브
doc.add_heading("7-5. 서비스별 주요 디자인 이니셔티브", level=3)

initiatives = [
    ("네이버 지도 V6 (2025.11)",
     "6년 만의 전면 리디자인. 단순 내비게이션 앱에서 생활 플랫폼으로 전환. "
     "상단 탭 5개(Discover·예약·대중교통·내비게이션·저장) 재편, 다국어 지원, 원탭 예약 검색 도입."),
    ("네이버페이 BI 리뉴얼 (2024)",
     "3D 아이콘 도입 및 핀테크 특화 브랜드 아이덴티티 재정립. Behance 공개."),
    ("CHZZK 브랜드",
     "라이브 스트리밍 플랫폼 고유 시각 언어 구축. 게임·엔터 Z세대 타깃 컬러·타이포 체계."),
    ("BAND 리브랜딩",
     "'커뮤니티 모티베이터' 컨셉으로 브랜드 아이덴티티 재정립. "
     "모임의 생애 주기 전반을 지원하는 감각적·경쾌한 디자인 방향성."),
    ("NAVER TechOne 사이니지 (iF 수상)",
     "약 10개 자회사가 입주한 공간을 위한 통합 웨이파인딩 시스템. "
     "직관적 동선 안내와 브랜드 일관성을 동시에 구현."),
]

for title, desc in initiatives:
    para = doc.add_paragraph(style="List Bullet")
    run_t = para.add_run(f"{title}: ")
    set_run_font(run_t, size=10, bold=True)
    run_d = para.add_run(desc)
    set_run_font(run_d, size=10)

doc.add_paragraph()

# 7-6. 디자인 철학 요약
doc.add_heading("7-6. 디자인 철학 요약", level=3)

philosophy = [
    ("일관성 (Consistency)",
     "189개 서비스 아이콘·NAVER 그린·타이포그래피를 하나의 시각 언어로 통합. "
     "어떤 플랫폼·환경에서도 '네이버다움'을 즉각 인식할 수 있도록 설계."),
    ("접근성 (Accessibility)",
     "나눔폰트 무료 배포, 다국어 UI 지원, 시각적 명암 대비 기준 준수. "
     "사용자 계층·기기 다양성을 포용하는 범용 디자인 지향."),
    ("기술과 인간의 조화 (Human-Tech Harmony)",
     "1784 사옥에서 구현된 로봇-인간 공존 공간 UX가 대표 사례. "
     "AI·로봇 기술을 인간 생활 동선에 자연스럽게 녹이는 경험 설계."),
    ("브랜드 확장성 (Scalability)",
     "국내 포털에서 글로벌 플랫폼(웹툰·KREAM·SNOW)으로 확장 시 "
     "각 서비스의 고유성을 유지하면서도 모기업 NAVER와의 연결성을 잃지 않는 패밀리 브랜드 전략."),
]

table_phil = doc.add_table(rows=len(philosophy), cols=2)
table_phil.style = "Table Grid"
for i, (k, v) in enumerate(philosophy):
    row = table_phil.rows[i]
    row.cells[0].text = k
    row.cells[1].text = v
    set_cell_bg(row.cells[0], "E8F5E9")
    for cell in row.cells:
        for para in cell.paragraphs:
            for run in para.runs:
                set_run_font(run, size=10)
    table_phil.columns[0].width = Cm(4.5)
    table_phil.columns[1].width = Cm(10.5)

doc.add_paragraph()

# 출처 안내
src = doc.add_paragraph(
    "※ 출처: navercorp.com/en/company/brandGuide · designcompass.org · "
    "ifdesign.com · fastcompany.com · logotyp.us · brandcolorcode.com"
)
for run in src.runs:
    set_run_font(run, size=9, color=(100, 100, 100))

doc.save(DOC_PATH)
print("추가 완료: 7. 디자인 컨셉 섹션이 naver_research.docx에 저장되었습니다.")
